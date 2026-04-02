"""
resume_extractor.py
-------------------
Extracts structured candidate details from resumes in any format.
Uses Ollama (llama3) running locally — no API key required.

Supports: PDF, DOCX, DOC, TXT, RTF, images (JPG/PNG/WEBP/GIF).

Usage:
    from resume_extractor import extract_resume

    result = extract_resume("path/to/resume.pdf")
    print(result)  # dict with structured fields

Dependencies:
    pip install ollama pypdf python-docx pillow
    # For scanned PDFs:
    pip install pdf2image && sudo apt install poppler-utils
    # For legacy .doc files:
    sudo apt install antiword

Ollama setup (if not already done):
    curl -fsSL https://ollama.com/install.sh | sh
    ollama pull phi4-mini
    ollama serve          # starts the local server on http://localhost:11434
"""

import base64
import json
import os
import subprocess
import tempfile
from pathlib import Path
from typing import Union

import ollama

# ---------------------------------------------------------------------------
# Config
# ---------------------------------------------------------------------------

OLLAMA_MODEL = "phi4-mini"                # change to "llama3.2-vision" for native image support
OLLAMA_HOST  = "http://localhost:11434"   # default Ollama address

# ---------------------------------------------------------------------------
# Text extraction helpers (one per format)
# ---------------------------------------------------------------------------

def _extract_text_pdf(path: str) -> str | None:
    """Extract text from a PDF. Returns None for scanned (image-only) PDFs."""
    try:
        from pypdf import PdfReader
        reader = PdfReader(path)
        pages = [page.extract_text() or "" for page in reader.pages]
        text = "\n".join(pages).strip()
        if len(text) > 100:
            return text
    except Exception:
        pass
    return None   # scanned PDF — caller will rasterise


def _extract_text_docx(path: str) -> str:
    """Extract text from a .docx file using python-docx."""
    try:
        from docx import Document
        doc = Document(path)
        parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(
                    cell.text.strip() for cell in row.cells if cell.text.strip()
                )
                if row_text:
                    parts.append(row_text)
        return "\n".join(parts)
    except Exception as e:
        raise RuntimeError(f"Failed to read DOCX: {e}")


def _extract_text_doc(path: str) -> str:
    """Convert legacy .doc to text via antiword or pandoc."""
    try:
        result = subprocess.run(
            ["antiword", path], capture_output=True, text=True, timeout=30
        )
        if result.returncode == 0 and result.stdout.strip():
            return result.stdout
    except FileNotFoundError:
        pass
    try:
        result = subprocess.run(
            ["pandoc", path, "-t", "plain"], capture_output=True, text=True, timeout=30
        )
        if result.returncode == 0:
            return result.stdout
    except FileNotFoundError:
        pass
    raise RuntimeError(
        "Cannot read .doc file. Install 'antiword' or 'pandoc':\n"
        "  sudo apt install antiword   OR   sudo apt install pandoc"
    )


def _extract_text_txt_rtf(path: str) -> str:
    """Read plain text or RTF (pandoc strips RTF markup)."""
    ext = Path(path).suffix.lower()
    if ext == ".rtf":
        try:
            result = subprocess.run(
                ["pandoc", path, "-t", "plain"], capture_output=True, text=True, timeout=30
            )
            if result.returncode == 0:
                return result.stdout
        except FileNotFoundError:
            pass
    with open(path, "r", errors="replace") as f:
        return f.read()


# ---------------------------------------------------------------------------
# Prompt
# ---------------------------------------------------------------------------

_SYSTEM_PROMPT = """You are a resume parser for a hiring platform.
Extract all important details from the provided resume text and return
ONLY a valid JSON object — no markdown fences, no extra commentary, nothing else.

The JSON must follow this schema exactly (use null for any missing field):
{
  "name": "string",
  "email": "string",
  "phone": "string",
  "location": "string",
  "linkedin": "string | null",
  "github": "string | null",
  "portfolio": "string | null",
  "summary": "string | null",
  "total_experience_years": "number | null",
  "current_role": "string | null",
  "skills": ["string"],
  "technical_skills": ["string"],
  "soft_skills": ["string"],
  "languages": ["string"],
  "work_experience": [
    {
      "title": "string",
      "company": "string",
      "location": "string | null",
      "start_date": "string | null",
      "end_date": "string | null",
      "is_current": "boolean",
      "responsibilities": ["string"]
    }
  ],
  "education": [
    {
      "degree": "string",
      "field": "string | null",
      "institution": "string",
      "location": "string | null",
      "graduation_year": "string | null",
      "gpa": "string | null"
    }
  ],
  "certifications": [
    {
      "name": "string",
      "issuer": "string | null",
      "year": "string | null"
    }
  ],
  "projects": [
    {
      "name": "string",
      "description": "string | null",
      "technologies": ["string"],
      "url": "string | null"
    }
  ],
  "awards": ["string"],
  "publications": ["string"],
  "volunteer_experience": ["string"]
}

Rules:
- Infer total_experience_years from work history dates when not stated explicitly.
- Split skills into technical_skills (tools, languages, frameworks) and soft_skills
  (communication, leadership, etc.) where possible; also keep all in "skills".
- Dates can be "MM/YYYY", "Month YYYY", "YYYY", or "Present".
- Return an empty array [] for list fields with no data, not null.
- Output ONLY the JSON. Do not add any text before or after it.
"""


# ---------------------------------------------------------------------------
# Ollama call helpers
# ---------------------------------------------------------------------------

def _make_ollama_client() -> ollama.Client:
    return ollama.Client(host=OLLAMA_HOST)


def _strip_json_fences(text: str) -> str:
    """Remove markdown code fences that some models add despite instructions."""
    text = text.strip()
    if text.startswith("```"):
        lines = text.splitlines()
        inner = lines[1:] if lines[-1].strip() == "```" else lines[1:]
        text = "\n".join(inner).rstrip("`").strip()
    return text


def _call_ollama_text(client: ollama.Client, resume_text: str) -> dict:
    """Send resume text to Ollama and return parsed JSON dict."""
    response = client.chat(
        model=OLLAMA_MODEL,
        messages=[
            {"role": "system", "content": _SYSTEM_PROMPT},
            {"role": "user",   "content": f"Resume:\n\n{resume_text}"},
        ],
        options={"temperature": 0},   # deterministic output
    )
    raw = response["message"]["content"].strip()
    raw = _strip_json_fences(raw)
    return json.loads(raw)


def _call_ollama_image(client: ollama.Client, image_path: str) -> dict:
    """
    Send an image to Ollama. Requires a vision-capable model (e.g. llama3.2-vision).
    Falls back to pytesseract OCR if the current model is text-only.
    """
    try:
        response = client.chat(
            model=OLLAMA_MODEL,
            messages=[
                {"role": "system", "content": _SYSTEM_PROMPT},
                {
                    "role": "user",
                    "content": "This is a resume image. Extract all details.",
                    "images": [image_path],   # Ollama accepts file paths
                },
            ],
            options={"temperature": 0},
        )
        raw = response["message"]["content"].strip()
        raw = _strip_json_fences(raw)
        return json.loads(raw)
    except Exception as vision_err:
        # Model doesn't support vision — fall back to pytesseract OCR
        try:
            import pytesseract
            from PIL import Image
            ocr_text = pytesseract.image_to_string(Image.open(image_path))
            if len(ocr_text.strip()) < 50:
                raise RuntimeError("OCR returned too little text.")
            return _call_ollama_text(client, ocr_text)
        except ImportError:
            raise RuntimeError(
                f"Model '{OLLAMA_MODEL}' does not support vision and pytesseract is not installed.\n"
                "To handle image resumes, either:\n"
                "  1. Use a vision model:  OLLAMA_MODEL = 'llama3.2-vision'\n"
                "     Then run: ollama pull llama3.2-vision\n"
                "  2. Install OCR fallback: pip install pytesseract pillow\n"
                "     Then run: sudo apt install tesseract-ocr\n"
                f"Original error: {vision_err}"
            )


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------

def extract_resume(
    file_path: Union[str, Path],
    model: str = None,
    ollama_host: str = None,
) -> dict:
    """
    Extract structured information from a resume file using a local Ollama model.

    Parameters
    ----------
    file_path : str | Path
        Path to the resume file. Supported formats:
        PDF, DOCX, DOC, TXT, RTF, JPG, JPEG, PNG, GIF, WEBP.
    model : str, optional
        Ollama model name. Defaults to OLLAMA_MODEL ('llama3').
        Use 'llama3.2-vision' for native image/scanned-PDF support.
    ollama_host : str, optional
        Ollama server URL. Defaults to OLLAMA_HOST ('http://localhost:11434').

    Returns
    -------
    dict
        Structured resume data. All schema keys are always present;
        missing values are None or [].

    Raises
    ------
    FileNotFoundError
        If the resume file does not exist.
    ValueError
        If the file format is not supported.
    RuntimeError
        If text or image extraction fails.
    ollama.ResponseError
        If the Ollama API call fails (e.g. model not pulled yet).
    json.JSONDecodeError
        If the model returns malformed JSON.
    """
    global OLLAMA_MODEL, OLLAMA_HOST
    if model:
        OLLAMA_MODEL = model
    if ollama_host:
        OLLAMA_HOST = ollama_host

    path = Path(file_path)
    if not path.exists():
        raise FileNotFoundError(f"File not found: {file_path}")

    client = _make_ollama_client()
    ext = path.suffix.lower()

    # --- Image formats ---
    if ext in {".jpg", ".jpeg", ".png", ".gif", ".webp"}:
        return _call_ollama_image(client, str(path))

    # --- Text-based formats ---
    if ext == ".pdf":
        text = _extract_text_pdf(str(path))
        if text is None:
            # Scanned PDF — rasterise first page then use vision/OCR
            try:
                from pdf2image import convert_from_path
                images = convert_from_path(str(path), first_page=1, last_page=1, dpi=200)
                with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as tmp:
                    images[0].save(tmp.name, "PNG")
                    tmp_path = tmp.name
                result = _call_ollama_image(client, tmp_path)
                os.unlink(tmp_path)
                return result
            except ImportError:
                raise RuntimeError(
                    "Scanned PDF detected. Install pdf2image + poppler-utils:\n"
                    "  pip install pdf2image\n"
                    "  sudo apt install poppler-utils"
                )
        return _call_ollama_text(client, text)

    if ext == ".docx":
        return _call_ollama_text(client, _extract_text_docx(str(path)))

    if ext == ".doc":
        return _call_ollama_text(client, _extract_text_doc(str(path)))

    if ext in {".txt", ".rtf", ".md"}:
        return _call_ollama_text(client, _extract_text_txt_rtf(str(path)))

    raise ValueError(
        f"Unsupported file format: '{ext}'. "
        "Supported: pdf, docx, doc, txt, rtf, jpg, jpeg, png, gif, webp."
    )


# ---------------------------------------------------------------------------
# CLI convenience
# ---------------------------------------------------------------------------

if __name__ == "__main__":
    import sys

    if len(sys.argv) < 2:
        print("Usage: python resume_extractor.py <resume_file> [output.json] [--model llama3.2-vision]")
        sys.exit(1)

    input_file = sys.argv[1]
    output_file = None
    model_override = None

    args = sys.argv[2:]
    i = 0
    while i < len(args):
        if args[i] == "--model" and i + 1 < len(args):
            model_override = args[i + 1]
            i += 2
        else:
            output_file = args[i]
            i += 1

    print(f"Extracting resume: {input_file}  (model: {model_override or OLLAMA_MODEL})")
    data = extract_resume(input_file, model=model_override)

    formatted = json.dumps(data, indent=2, ensure_ascii=False)

    if output_file:
        with open(output_file, "w") as f:
            f.write(formatted)
        print(f"Saved to {output_file}")
    else:
        print(formatted)